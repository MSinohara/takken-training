from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from pathlib import Path


OUTPUT = "docs/事務局用操作説明書_第1版.docx"
SCREEN_DIR = Path("docs/manual_screens")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D7DCE3", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn("w:" + margin_name))
        if node is None:
            node = OxmlElement("w:" + margin_name)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Yu Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Yu Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Yu Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("研修会参加確認システム")
    set_run_font(run, 22, True, "0B2545")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("事務局用 操作説明書（第1版）")
    set_run_font(run, 18, True, "2E74B5")

    table = doc.add_table(rows=4, cols=2)
    set_table_width(table, [2400, 6960])
    rows = [
        ("対象", "事務局担当者、研修会運営担当者"),
        ("目的", "研修会の作成、案内、受付、履歴確認、集計、修了証発行までの基本操作を確認するための説明書です。"),
        ("版", "第1版"),
        ("作成日", "2026年7月"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        set_cell_shading(row.cells[0], "E8EEF5")
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, 10.5, label == paragraph.text, "000000")

    add_callout(
        doc,
        "第1版の位置づけ",
        "本書は、事務局が全体像をつかみ、通常業務で迷わないことを目的にした操作説明書です。"
        "研修会の作成、案内、受付、履歴確認、集計、修了証発行までの基本的な流れを確認できます。"
    )


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    set_cell_border(cell, "C8D2DE")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, 10.5, True, "1F3A5F")
    p = cell.add_paragraph(body)
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        set_run_font(run, 10.5)


def add_screenshot_placeholder(doc, caption):
    return


def add_capture(doc, filename, caption, fallback):
    path = SCREEN_DIR / filename
    if not path.exists():
        return

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("画面例：" + caption)
    set_run_font(r, 9.5, True, "1F4D78")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    p.add_run().add_picture(str(path), width=Inches(6.1))


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.38)
        p.paragraph_format.first_line_indent = Inches(-0.19)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)
        for run in p.runs:
            set_run_font(run, 10.5)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.38)
        p.paragraph_format.first_line_indent = Inches(-0.19)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)
        for run in p.runs:
            set_run_font(run, 10.5)


def add_simple_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, "E8EEF5")
        cell.text = header
        for p in cell.paragraphs:
            for r in p.runs:
                set_run_font(r, 9.5, True, "0B2545")
    for row_data in rows:
        row = table.add_row()
        for cell, value in zip(row.cells, row_data):
            cell.text = value
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    set_run_font(r, 9.5)
    return table


def add_section_flow(doc):
    doc.add_heading("1. システム全体の流れ", level=1)
    add_callout(
        doc,
        "最初に押さえる考え方",
        "このシステムは、研修会の作成から、案内メール、当日受付、参加履歴、集計、修了証発行までを一連で管理します。"
        "細かい設定は多くありますが、通常業務では下記の順番を押さえれば運用できます。"
    )
    add_numbers(doc, [
        "会員マスタ・組織設定を準備する。",
        "研修会を作成し、対象者・受付方式・メール内容を設定する。",
        "必要に応じて案内メールを送信する。",
        "当日はQR受付、検索受付、係員受付などで受付する。",
        "受付モニターと参加履歴で受付状況を確認する。",
        "研修会別・年度別に集計する。",
        "必要に応じて修了証対象者を抽出し、修了証PDFを発行する。",
    ])


def add_login_menu(doc):
    doc.add_heading("2. ログインとメニュー", level=1)
    doc.add_paragraph(
        "ログイン管理を有効にしている場合は、ユーザーIDとパスワードでログインしてから操作します。"
        "担当者の権限により、表示されるメニューや操作できる内容が変わります。"
    )
    add_simple_table(
        doc,
        ["メニュー", "主な用途"],
        [
            ("研修会管理", "研修会一覧、研修会作成、メール送信、受付、参加履歴、集計を行います。"),
            ("会員・組織管理", "会員マスタ、所属組織、個人会員、役員・来賓などを管理します。"),
            ("修了証管理", "修了証ルール、発行者、対象者抽出、PDF発行、発行履歴を管理します。"),
            ("システム設定", "ログイン権限など、システム全体の設定を管理します。"),
        ],
        [2400, 6960],
    )
    add_capture(
        doc,
        "01_top.png",
        "トップ画面",
        "システム全体の入口",
    )
    add_capture(
        doc,
        "01b_training_menu.png",
        "研修会管理メニュー",
        "研修会一覧、メール送信、受付、参加履歴、集計への入口",
    )


def add_member_org(doc):
    doc.add_heading("3. 会員マスタ・組織設定", level=1)
    doc.add_heading("3.1 会員マスタの取り込み", level=2)
    doc.add_paragraph(
        "会員マスタは、受付対象者や検索受付、集計、メール送信の基礎データです。"
        "CSV取り込み時は、列名で判定するため、列の順番よりも見出し名が重要です。"
    )
    add_simple_table(
        doc,
        ["項目", "説明"],
        [
            ("業者番号", "会員を識別する番号です。QR受付や検索受付の基準になります。"),
            ("会社名", "受付画面、モニター、履歴、メール宛名などに表示されます。"),
            ("代表者名", "会社単位受付の代表者、個人会員の代表者自動登録などに利用できます。"),
            ("ブロック・支部・地区", "対象者の絞り込み、モニター表示、集計に利用します。"),
            ("メール", "案内メールの送信先です。未入力でも取り込み自体は可能ですが、メール送信対象にはなりません。"),
        ],
        [2200, 7160],
    )
    add_callout(
        doc,
        "取り込み前の確認",
        "必須項目に空欄がある場合は、取り込み前に確認してください。特にブロック・支部・地区は、対象者抽出や集計に関わるため、原則入力しておきます。"
    )
    doc.add_heading("3.2 組織設定", level=2)
    add_bullets(doc, [
        "青年部、レディス部、研修委員会など、支部・地区とは別の所属を管理できます。",
        "会社に組織を付けると、代表者扱いの個人会員にも同じ所属を持たせる運用ができます。",
        "社員など2人目以降の個人会員は、必要に応じて個別に所属を設定します。",
    ])
    doc.add_heading("3.3 個人会員", level=2)
    doc.add_paragraph(
        "研修会ごとに受付単位を選択できます。会社単位受付では1社を1受付として扱い、個人単位受付では代表者や社員など個人ごとのQRで受付できます。"
        "青年部、レディス部、委員会など、個人単位での参加状況を確認したい研修で利用します。"
    )
    add_capture(
        doc,
        "10_member_menu.png",
        "会員・組織管理メニュー",
        "会員一覧、会員詳細、組織設定、個人会員マスタ",
    )


def add_training(doc):
    doc.add_heading("4. 研修会の作成・編集", level=1)
    doc.add_paragraph(
        "研修会作成では、研修名、開催日、主催区分、対象者、受付方式、メール内容などを登録します。"
        "研修会一覧から対象の研修会を開くと、研修会詳細画面で受付や集計に進めます。"
    )
    add_simple_table(
        doc,
        ["設定項目", "ポイント"],
        [
            ("開催日・年度", "一覧や集計の絞り込みに利用します。"),
            ("主催区分", "ブロック研修、支部研修などの分類です。"),
            ("対象ブロック・支部・地区", "受付対象者や集計の基準になります。"),
            ("対象組織", "青年部、レディス部、研修委員会などの研修で利用します。"),
            ("受付方式", "会社QR、会場QR＋検索、スマホ登録、個人単位受付など、研修会に合わせて選択します。"),
            ("会場", "会場マスタと紐づけることで、位置情報受付などにも活用できます。"),
            ("添付PDF", "案内メールに添付する資料です。差し替え・削除時はDrive上の扱いにも注意します。"),
        ],
        [2600, 6760],
    )
    add_callout(
        doc,
        "運用上の注意",
        "研修会の受付方式や対象者設定は、案内メール送信前に確認してください。"
        "当日直前に変更する場合は、メール本文やQR案内との整合性も確認します。"
    )
    doc.add_heading("4.1 研修会作成時の入力例", level=2)
    doc.add_paragraph(
        "新しい研修会を作る場合は、まず基本情報を入れ、その後に対象者、受付方式、メール内容を確認します。"
        "下記はブロック研修を作る場合の入力例です。"
    )
    add_simple_table(
        doc,
        ["入力項目", "入力例", "説明"],
        [
            ("研修名", "第1回 第十ブロック研修会", "参加者に表示される名称です。案内メールや履歴にも表示されます。"),
            ("開催日", "2026/07/14", "年度別一覧や集計の基準になります。"),
            ("主催区分", "ブロック研修", "ブロック研修、支部研修などの分類です。"),
            ("対象ブロック", "第十ブロック", "第十ブロック所属会員を対象にする場合に設定します。"),
            ("対象支部", "なし", "ブロック全体を対象にする場合は支部を指定しません。杉並支部研修などでは支部を指定します。"),
            ("対象地区", "なし", "地区単位で絞る場合のみ指定します。"),
            ("対象組織", "なし", "青年部、レディス部、研修委員会など組織限定の研修で指定します。"),
            ("受付単位", "会社単位", "1社を1受付として扱う場合は会社単位、個人ごとに受付する場合は個人単位を選びます。"),
            ("受付方式", "会場QR＋検索受付", "当日、会場QRから会社名または業者番号で検索して受付する方式です。"),
            ("位置情報受付", "必要に応じて使用", "他の受付方法に追加できる補助受付です。使わない場合は無効にします。"),
            ("会場", "登録済み会場を選択", "位置情報受付を使う場合は、会場の緯度経度や受付可能距離も確認します。"),
            ("メール件名", "第1回 第十ブロック研修会のお知らせ", "送信前に必ずプレビューで確認します。"),
            ("メール本文", "日時、会場、当日の受付方法を記載", "受付方式に応じた案内文やURLが追加されます。"),
        ],
        [1700, 2500, 5160],
    )
    add_callout(
        doc,
        "支部研修の入力例",
        "杉並支部研修の場合は、対象ブロックに第十ブロック、対象支部に杉並支部を指定します。"
        "中野支部、世田谷支部も同様に、対象支部を変えて作成します。"
    )
    add_capture(
        doc,
        "02_training_list.png",
        "研修会一覧",
        "年度、開催状態、修了証有無で研修会を探します。",
    )
    add_capture(
        doc,
        "03_training_detail.png",
        "研修会詳細",
        "メール送信、受付、受付モニター、参加履歴、集計へ進む起点です。",
    )


def add_mail(doc):
    doc.add_heading("5. メール送信", level=1)
    doc.add_paragraph(
        "研修会詳細またはメール送信画面から、対象者へ案内メールを送信できます。"
        "メール本文には、受付方式に応じた案内文やURLが入ります。"
    )
    add_simple_table(
        doc,
        ["受付方式", "メール本文に入る主な案内"],
        [
            ("会社QR受付", "会社ごとの受付用QRリンクを案内します。"),
            ("会場QR＋検索受付", "当日、会場に掲示されたQRから会社名または業者番号で受付する旨を案内します。"),
            ("スマホ登録＋会場QR受付", "事前にスマホへ会社情報を登録し、当日会場QRから受付する旨を案内します。"),
            ("位置情報受付", "会社QR、会場QR＋検索、スマホ登録、個人単位受付などに追加して利用できる補助的な受付リンクを案内します。"),
            ("個人単位受付", "個人QRのリンクを案内します。"),
        ],
        [2600, 6760],
    )
    add_bullets(doc, [
        "送信前に、件名、本文、署名、添付PDFを確認します。",
        "必要に応じてメールプレビューで、実際に届く文面を確認します。",
        "一括送信前は送信対象者数を確認し、誤送信を防ぎます。",
        "送信後はメール送信履歴で結果を確認できます。",
    ])
    add_capture(
        doc,
        "04_mail_send.png",
        "メール送信画面",
        "送信対象、件名、本文、プレビューを確認してから送信します。",
    )


def add_reception(doc):
    doc.add_heading("6. 当日受付", level=1)
    doc.add_paragraph(
        "当日は、研修会詳細画面から受付画面や受付モニターへ進みます。"
        "受付方式は研修会ごとに異なりますが、受付結果は参加履歴に記録されます。"
    )
    add_simple_table(
        doc,
        ["受付方法", "概要"],
        [
            ("会社QR受付", "事前に配布した会社別QRを、受付担当者が読み取ります。"),
            ("会場QR＋検索受付", "参加者が会場QRを読み込み、会社名や業者番号で検索して受付します。"),
            ("スマホ登録＋会場QR受付", "事前登録済みのスマホから、会場QRを読み込んで受付します。"),
            ("位置情報受付", "他の受付方法と併用できる補助的な受付です。会場付近にいることを確認して受付します。"),
            ("係員受付", "受付担当者が会社名・業者番号などで検索して受付します。"),
            ("一般・他ブロック・役員来賓", "会員マスタ対象外の参加者も、参考参加者として受付できます。"),
        ],
        [2600, 6760],
    )
    add_callout(
        doc,
        "誤受付時の考え方",
        "間違えて受付した場合は、参加履歴または受付管理画面から取消を行います。取消理由を選択し、必要に応じて復活できます。"
    )
    add_capture(
        doc,
        "05_live_checkin.png",
        "受付モニター",
        "当日の受付状況、対象人数、未受付者などを確認します。",
    )
    add_capture(
        doc,
        "06_staff_checkin.png",
        "係員用受付",
        "第十ブロック会員、他ブロック会員、一般参加者、当日参加予定者を係員が受付できます。",
    )


def add_history_stats(doc):
    doc.add_heading("7. 参加履歴・集計", level=1)
    doc.add_heading("7.1 参加履歴", level=2)
    doc.add_paragraph(
        "参加履歴では、誰が、いつ、どの方法で受付したかを確認できます。"
        "取消・復活の情報も確認できるため、当日の問い合わせ対応や後日の確認に利用します。"
    )
    doc.add_heading("7.2 研修会ごとの集計", level=2)
    add_bullets(doc, [
        "研修会詳細の集計は、その研修会の現在の参加状況を確認するためのものです。",
        "対象者数、受付済み数、参加率、支部・地区・組織別の参加状況を確認できます。",
        "一般参加者、他ブロック会員、役員来賓は参加率の対象外として参考件数に分けて表示します。",
    ])
    doc.add_heading("7.3 年間集計", level=2)
    doc.add_paragraph(
        "年間集計では、年度内の研修会を横断して、支部別・地区別・組織別などの参加傾向を確認します。"
        "参加フォロー分析と組み合わせることで、参加が少ない会員の確認にも利用できます。"
    )
    add_capture(
        doc,
        "07_history.png",
        "参加履歴",
        "受付済み、取消済み、復活済みなどの履歴を確認します。",
    )
    add_capture(
        doc,
        "08_annual_stats.png",
        "年間研修会集計",
        "年度内の研修会をまとめて、参加状況を横断的に確認します。",
    )


def add_certificate(doc):
    doc.add_heading("8. 修了証", level=1)
    doc.add_paragraph(
        "修了証機能では、指定した条件に合う参加者を抽出し、修了証PDFを発行できます。"
        "条件や発行者は事前に登録し、発行履歴で誰にどの修了証を発行したかを確認します。"
    )
    add_simple_table(
        doc,
        ["機能", "説明"],
        [
            ("修了証ルール", "対象年度、対象支部・組織、必要参加回数、必須研修などを設定します。"),
            ("発行者設定", "発行者の肩書、役職、氏名を登録し、ルールから選択できます。"),
            ("対象者抽出", "ルールに合う対象者を修了証対象として抽出します。"),
            ("PDF発行", "修了証テンプレートに会社名、参加研修、発行者などを差し込みPDF化します。"),
            ("発行履歴", "発行済みPDF、発行日時、送信先などを確認します。"),
        ],
        [2600, 6760],
    )
    add_callout(
        doc,
        "運用上の注意",
        "修了証発行画面は、これから発行する対象者を確認する画面です。発行後の確認は修了証履歴で行います。"
    )
    add_capture(
        doc,
        "09_certificate_menu.png",
        "修了証管理メニュー",
        "修了証ルール、修了証対象、修了証履歴、PDFプレビュー",
    )


def add_troubleshooting(doc):
    doc.add_heading("9. よくある対応", level=1)
    add_simple_table(
        doc,
        ["困ったこと", "確認すること"],
        [
            ("メールが届かない", "メールアドレス、送信対象、送信履歴、迷惑メールを確認します。"),
            ("QRが読めない", "画面の明るさ、QRの種類、受付方式、対象研修会を確認します。"),
            ("会社名が見つからない", "会員マスタ、対象支部・地区、表記ゆれを確認します。"),
            ("受付モニターに出ない", "受付が完了しているか、表示フィルタ、対象研修会を確認します。"),
            ("間違えて受付した", "受付取消を行い、必要であれば復活します。"),
            ("添付PDFを差し替えたい", "研修会編集で添付PDFを削除し、新しいPDFを登録します。"),
            ("修了証対象に出ない", "ルールID、対象年度、対象支部・組織、参加履歴、発行済みかを確認します。"),
        ],
        [2500, 6860],
    )


def add_appendix(doc):
    doc.add_heading("10. 説明時に確認する主な画面", level=1)
    doc.add_paragraph(
        "操作説明では、次の順番で画面を確認すると、全体の流れを説明しやすくなります。"
    )
    add_simple_table(
        doc,
        ["重要度", "画面", "説明する内容"],
        [
            ("高", "トップ画面・研修会管理メニュー", "全体の入口と、研修会管理へ進む流れを確認します。"),
            ("高", "研修会一覧・研修会詳細", "研修会ごとの操作の起点、メール送信、受付、履歴、集計への進み方を確認します。"),
            ("高", "受付モニター", "当日の受付状況、対象人数、未受付者の確認方法を確認します。"),
            ("高", "係員受付・会場QR受付", "受付担当者が実際に使う受付方法を確認します。"),
            ("中", "メール送信・メールプレビュー", "案内メールの確認、テスト送信、一括送信の流れを確認します。"),
            ("中", "参加履歴・集計", "受付後に残る記録と、研修会別・年度別の集計を確認します。"),
            ("中", "修了証管理", "修了証ルール、対象者抽出、PDF発行、発行履歴を確認します。"),
            ("低", "詳細設定・各種モーダル", "必要に応じて、組織設定、権限設定、修了証ルール詳細などを確認します。"),
        ],
        [1200, 3000, 5160],
    )
    add_callout(
        doc,
        "説明の進め方",
        "最初に全体メニューを見せ、次に1つの研修会を開いて、メール送信、当日受付、受付モニター、参加履歴、集計の順に説明すると流れが伝わりやすくなります。"
    )


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("研修会参加確認システム 事務局用操作説明書 第1版")
    set_run_font(r, 8.5, False, "666666")


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_footer(section)

    setup_styles(doc)
    add_title(doc)
    doc.add_page_break()

    add_section_flow(doc)
    add_login_menu(doc)
    add_member_org(doc)
    add_training(doc)
    add_mail(doc)
    add_reception(doc)
    add_history_stats(doc)
    add_certificate(doc)
    add_troubleshooting(doc)
    add_appendix(doc)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
